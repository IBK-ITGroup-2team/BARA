from mimetypes import encodings_map
from msilib.schema import MIME
from telnetlib import DO
from docx import Document
from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.shared import Cm, Inches
from docx.text.run import Font
from docx.oxml.ns import qn

import smtplib
import os
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication  # 메일의 첨부 파일을 base64 형식으로 변환
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email import encoders

import csv

import WordCloud
# import IndividualCrawling


def main():
    # 워드클라우드의 main() 함수가 선행되어야함! 지금은 크롤링코드에서 실행연결 되어있음.
    # WordCloud.TOSS_negative_top3 이런 식으로 가져와서 쓰면 됨.
    #print(WordCloud.TOSS_negative_top3)
    #print(WordCloud.TOSS_positive_top3)

    # best 은행 선정
    
    # 개인 고객용 어플
    starListIndividual=[]
    bestBankI=''
    
    # 하나 은행
    fHANA=open('./reviews/별점/HANAreview_rating.csv','r',encoding='utf-8')
    rdrHANA=csv.reader(fHANA)
    for line in rdrHANA:
        starListIndividual.append(line[1])
    fHANA.close()

    # KB
    fKB=open('./reviews/별점/KBreview_rating.csv','r',encoding='utf-8')
    rdrKB=csv.reader(fKB)
    for line in rdrKB:
        starListIndividual.append(line[1])
    fKB.close()     
    
    # WOORI
    fWOORI=open('./reviews/별점/WONreview_rating.csv','r',encoding='utf-8')
    rdrWOORI=csv.reader(fWOORI)
    for line in rdrWOORI:
        starListIndividual.append(line[1])
    fWOORI.close()  
    
    # NH
    fNH=open('./reviews/별점/NHreview_rating.csv','r',encoding='utf-8')
    rdrNH=csv.reader(fNH)
    for line in rdrNH:
        starListIndividual.append(line[1])
    fNH.close() 
    
    # SHINHAN
    fSH=open('./reviews/별점/신한review_rating.csv','r',encoding='utf-8')
    rdrSH=csv.reader(fSH)
    for line in rdrSH:
        starListIndividual.append(line[1])
    fSH.close()
    
    bestScoreI=max(starListIndividual)
    
    indexI=starListIndividual.index(bestScoreI)
    if indexI == 2:
        bestWordI=WordCloud.SHINHAN_positive_top3[0]
        bestBankI='신한은행'
        
    elif indexI == 4:
        bestWordI=WordCloud.KB_positive_top3[0]
        bestBankI='국민은행'
        
    elif indexI == 6:
        bestWordI=WordCloud.WOORI_positive_top3[0]
        bestBankI='우리은행'
        
    elif indexI ==8:
        bestWordI=WordCloud.NH_positive_top3[0]
        bestBankI='농협은행'
        
    else:
        bestWordI=WordCloud.HANA_positive_top3[0]
        bestBankI='하나은행'
        
        
    #기업용 고객 어플 베스트 선정
    starListEnterprise=[]
    bestBankE=''
    
    # 하나 은행
    fHANAE=open('./reviews/별점/HANA_enterprise_review_rating.csv','r',encoding='utf-8')
    rdrHANAE=csv.reader(fHANAE)
    for line in rdrHANAE:
        starListEnterprise.append(line[1])
    fHANAE.close()

    # KB
    fKBE=open('./reviews/별점/KBreview_rating(KB스타기업뱅킹).csv','r',encoding='utf-8')
    rdrKBE=csv.reader(fKBE)
    for line in rdrKBE:
        starListEnterprise.append(line[1])
    fKB.close()     
    
    # WOORI
    fWOORIE=open('./reviews/별점/WOORIbank_enterprise_review_rating.csv','r',encoding='utf-8')
    rdrWOORIE=csv.reader(fWOORIE)
    for line in rdrWOORIE:
        starListEnterprise.append(line[1])
    fWOORIE.close()  
    
    # NH
    fNHE=open('./reviews/별점/NHreview_rating(NH기업뱅킹).csv','r',encoding='utf-8')
    rdrNHE=csv.reader(fNHE)
    for line in rdrNHE:
        starListEnterprise.append(line[1])
    fNHE.close() 
    
    # SHINHAN
    fSHE=open('./reviews/별점/SHINHANbank_enterprise_review_rating.csv','r',encoding='utf-8')
    rdrSHE=csv.reader(fSHE)
    for line in rdrSHE:
        starListEnterprise.append(line[1])
    fSHE.close()
    
    bestScoreE=max(starListEnterprise)
    
    indexE=starListEnterprise.index(bestScoreE)
    if indexE == 2:
        bestWordE=WordCloud.HANA_E_positive_top3[0]
        bestBankE='하나은행'
        
    elif indexE == 4:
        bestWordE=WordCloud.KB_E_positive_top3[0]
        bestBankE='국민은행'
        
    elif indexE == 6:
        bestWordE=WordCloud.SHINHAN_E_positive_top3[0]
        bestBankE='신한은행'
        
    elif indexE ==8:
        bestWordE=WordCloud.NH_E_positive_top3[0]
        bestBankE='농협은행'
        
    else:
        bestWordE=WordCloud.WOORI_E_positive_top3[0]
        bestBankE='우리은행'
        
    # 인터넷 전문 은행 베스트 선정
    starListInternet=[]
    bestBankInternet=''
    
     # TOSS
    fTOSS=open('./reviews/별점/TOSS_review_rating.csv','r',encoding='utf-8')
    rdrTOSS=csv.reader(fTOSS)
    for line in rdrTOSS:
        starListInternet.append(line[1])
    fTOSS.close()  
    
    # KAKAO
    fKAKAO=open('./reviews/별점/KAKAO_review_rating.csv','r',encoding='utf-8')
    rdrKAKAO=csv.reader(fKAKAO)
    for line in rdrKAKAO:
        starListInternet.append(line[1])
    fKAKAO.close() 
    
    # Kbank
    fKbank=open('./reviews/별점/KBank_review_rating.csv','r',encoding='utf-8')
    rdrKbank=csv.reader(fKbank)
    for line in rdrKbank:
        starListInternet.append(line[1])
    fKbank.close()
    
    bestScoreInternet=max(starListInternet)
    
    indexInternet=starListInternet.index(bestScoreInternet)
    if indexInternet == 2:
        bestWordInternet=WordCloud.TOSS_positive_top3[0]
        bestBankInternet='토스'
        
    elif indexInternet == 4:
        bestWordInternet=WordCloud.KBank_positive_top3[0]
        bestBankInternet='케이뱅크'
        
    else:
        bestWordInternet=WordCloud.KAKAO_positive_top3[0]
        bestBankInternet='카카오뱅크'

        
    document = Document()
    # 스타일 적용
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'

    banks = ['하나', '우리', '신한', '국민', '농협']
    internetBanks = ['카카오뱅크', '케이뱅크', '토스']

    # 맨 윗 줄
    document.add_picture('sender.PNG',width=Cm(16),height=Cm(1))
    
    # 제목
    #title=document.add_paragraph('')
    document.add_picture("title.PNG", width=Cm(16), height=Cm(2))
    
    document.add_paragraph('')
    dateToday = datetime.today()
    document.add_paragraph(datetime.today().strftime("%Y. %m. %d"))  # 해당 날짜
    
    

    document.add_paragraph('')
    document.add_paragraph('')

    objective = document.add_paragraph('')
    objective.add_run('□ 당행과 타행의 개인고객용 모바일 앱 사용자 반응 분석').bold = True
    document.add_paragraph('    ○ 당행 개인고객용 모바일 앱 (i-one bank) 사용자 반응 분석')

    # 표 생성
    grid_t_style = document.styles["Table Grid"]
    IBKTable = document.add_table(3, 2, grid_t_style)

    IBKTableCells1 = IBKTable.rows[0].cells
    IBKTableCells1[0].paragraphs[0].add_run('긍정적 반응')
    IBKTableCells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # IBKTableCells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    IBKTableCells1 = IBKTable.rows[0].cells
    IBKTableCells1[1].paragraphs[0].add_run('부정적 반응')
    IBKTableCells1[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # IBKTableCells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    IBKCell10 = IBKTable.cell(1, 0)
    IBKPara10 = IBKCell10.add_paragraph()
    IBKRun10 = IBKPara10.add_run()
    IBKRun10.add_picture("./wordcloud/개인고객/IBK_WordCloud_P.png", width=Cm(7), height=Cm(5))

    IBKCell11 = IBKTable.cell(1, 1)
    IBKPara11 = IBKCell11.add_paragraph()
    IBKRun11 = IBKPara11.add_run()
    IBKRun11.add_picture("./wordcloud/개인고객/IBK_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    IBKTableCells3 = IBKTable.rows[2].cells
    IBKTableCells3[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    IBKTableCells3[0].paragraphs[0].add_run('1. ' + WordCloud.IBK_positive_top3[0] + '\n')
    IBKTableCells3[0].paragraphs[0].add_run('2. ' + WordCloud.IBK_positive_top3[1] + '\n')
    IBKTableCells3[0].paragraphs[0].add_run('3. ' + WordCloud.IBK_positive_top3[2] + '\n')

    # 부정 빈출 단어 Top3
    IBKTableCells3[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    IBKTableCells3[1].paragraphs[0].add_run('1. '+ WordCloud.IBK_negative_top3[0]+'\n')
    IBKTableCells3[1].paragraphs[0].add_run('2. '+WordCloud.IBK_negative_top3[1]+'\n')
    IBKTableCells3[1].paragraphs[0].add_run('3. '+WordCloud.IBK_negative_top3[2]+'\n')

    for i in range(6):
        document.add_paragraph('')

    #  하나은행 개인 앱 리뷰 현황
    document.add_paragraph('    ○ 타행 개인고객용 모바일 앱 사용자 반응 분석').bold = True
    document.add_paragraph('        - 하나은행')

    # 하나은행 표 생성
    HANA1Table = document.add_table(3, 2, grid_t_style)

    HANA1Cells1 = HANA1Table.rows[0].cells
    HANA1Cells1[0].paragraphs[0].add_run('긍정적 반응').bold = True
    HANA1Cells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # HANA1Cells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    HANA1TableCells1 = HANA1Table.rows[0].cells
    HANA1Cells1[1].paragraphs[0].add_run('부정적 반응').bold = True
    HANA1TableCells1[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # HANA1TableCells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    HANACell10 = HANA1Table.cell(1, 0)
    HANAPara10 = HANACell10.add_paragraph()
    HANARun10 = HANAPara10.add_run()
    HANARun10.add_picture("./wordcloud/개인고객/HANA_WordCloud_P.png", width=Cm(7), height=Cm(5))

    HANACell11 = HANA1Table.cell(1, 1)
    HANAPara11 = HANACell11.add_paragraph()
    HANARun11 = HANAPara11.add_run()
    HANARun11.add_picture("./wordcloud/개인고객/HANA_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    HANATableCells3 = HANA1Table.rows[2].cells
    HANATableCells3[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    HANATableCells3[0].paragraphs[0].add_run('1. '+WordCloud.HANA_positive_top3[0]+'\n')
    HANATableCells3[0].paragraphs[0].add_run('2. '+WordCloud.HANA_positive_top3[1]+'\n')
    HANATableCells3[0].paragraphs[0].add_run('3. '+WordCloud.HANA_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    HANATableCells3[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    HANATableCells3[1].paragraphs[0].add_run('1. '+WordCloud.HANA_negative_top3[0]+'\n')
    HANATableCells3[1].paragraphs[0].add_run('2. '+WordCloud.HANA_negative_top3[1]+'\n')
    HANATableCells3[1].paragraphs[0].add_run('3. '+WordCloud.HANA_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')

    # 국민은행 앱 리뷰 현황
    document.add_paragraph('        - 국민은행')

    # 국민은행 표 생성
    KB1Table = document.add_table(3, 2, grid_t_style)

    KBCells1 = KB1Table.rows[0].cells
    KBCells1[0].paragraphs[0].add_run('긍정적 반응').bold = True
    KBCells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # KBCells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    KB1TableCells1 = KB1Table.rows[0].cells
    KBCells1[1].paragraphs[0].add_run('부정적 반응').bold = True
    KB1TableCells1[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # KB1TableCells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    KBCell10 = KB1Table.cell(1, 0)
    KBPara10 = KBCell10.add_paragraph()
    KBRun10 = KBPara10.add_run()
    KBRun10.add_picture("./wordcloud/개인고객/KB_WordCloud_P.png", width=Cm(7), height=Cm(5))

    KBCell11 = KB1Table.cell(1, 1)
    KBPara11 = KBCell11.add_paragraph()
    KBRun11 = KBPara11.add_run()
    KBRun11.add_picture("./wordcloud/개인고객/KB_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    KBTableCells3 = KB1Table.rows[2].cells
    KBTableCells3[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    KBTableCells3[0].paragraphs[0].add_run('1. '+WordCloud.KB_positive_top3[0]+'\n')
    KBTableCells3[0].paragraphs[0].add_run('2. '+WordCloud.KB_positive_top3[1]+'\n')
    KBTableCells3[0].paragraphs[0].add_run('3. '+WordCloud.KB_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    KBTableCells3[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    KBTableCells3[1].paragraphs[0].add_run('1. '+WordCloud.KB_negative_top3[0]+'\n')
    KBTableCells3[1].paragraphs[0].add_run('2. '+WordCloud.KB_negative_top3[1]+'\n')
    KBTableCells3[1].paragraphs[0].add_run('3. '+WordCloud.KB_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')

    # 신한은행 앱 리뷰 현황
    document.add_paragraph('        - 신한은행')

    # 신한은행 표 생성
    SH1Table = document.add_table(3, 2, grid_t_style)

    SHCells1 = SH1Table.rows[0].cells
    SHCells1[0].paragraphs[0].add_run('긍정적 반응').bold = True
    SHCells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # SHCells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    SH1TableCells1 = SH1Table.rows[0].cells
    SHCells1[1].paragraphs[0].add_run('부정적 반응').bold = True
    SH1TableCells1[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # SH1TableCells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입-신한
    SHCell10 = SH1Table.cell(1, 0)
    SHPara10 = SHCell10.add_paragraph()
    SHRun10 = SHPara10.add_run()
    SHRun10.add_picture("./wordcloud/개인고객/SHINHAN_WordCloud_P.png", width=Cm(7), height=Cm(5))

    SHCell11 = SH1Table.cell(1, 1)
    SHPara11 = SHCell11.add_paragraph()
    SHRun11 = SHPara11.add_run()
    SHRun11.add_picture("./wordcloud/개인고객/SHINHAN_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    SHTableCells3 = SH1Table.rows[2].cells
    SHTableCells3[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    SHTableCells3[0].paragraphs[0].add_run('1. '+WordCloud.SHINHAN_positive_top3[0]+'\n')
    SHTableCells3[0].paragraphs[0].add_run('2. '+WordCloud.SHINHAN_positive_top3[1]+'\n')
    SHTableCells3[0].paragraphs[0].add_run('3. '+WordCloud.SHINHAN_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    SHTableCells3[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    SHTableCells3[1].paragraphs[0].add_run('1. '+WordCloud.SHINHAN_negative_top3[0]+'\n')
    SHTableCells3[1].paragraphs[0].add_run('2. '+WordCloud.SHINHAN_negative_top3[1]+'\n')
    SHTableCells3[1].paragraphs[0].add_run('3. '+WordCloud.SHINHAN_negative_top3[2]+'\n')

    for i in range(3):
        document.add_paragraph('')

    # 농협 은행 개인 앱 리뷰 현황
    document.add_paragraph('        - 농협은행')

    # 농협은행 표 생성
    NH1Table = document.add_table(3, 2, grid_t_style)

    NHCells1 = NH1Table.rows[0].cells
    NHCells1[0].paragraphs[0].add_run('긍정적 반응').bold = True
    NHCells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # NHCells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    NH1TableCells1 = NH1Table.rows[0].cells
    NHCells1[1].paragraphs[0].add_run('부정적 반응').bold = True
    NH1TableCells1[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # NH1TableCells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입-농협
    NHCell10 = NH1Table.cell(1, 0)
    NHPara10 = NHCell10.add_paragraph()
    NHRun10 = NHPara10.add_run()
    NHRun10.add_picture("./wordcloud/개인고객/NH_WordCloud_P.png", width=Cm(7), height=Cm(5))

    NHCell11 = NH1Table.cell(1, 1)
    NHPara11 = NHCell11.add_paragraph()
    NHRun11 = NHPara11.add_run()
    NHRun11.add_picture("./wordcloud/개인고객/NH_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    NHTableCells3 = NH1Table.rows[2].cells
    NHTableCells3[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    NHTableCells3[0].paragraphs[0].add_run('1. '+WordCloud.NH_positive_top3[0]+'\n')
    NHTableCells3[0].paragraphs[0].add_run('2. '+WordCloud.NH_positive_top3[1]+'\n')
    NHTableCells3[0].paragraphs[0].add_run('3. '+WordCloud.NH_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    NHTableCells3[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    NHTableCells3[1].paragraphs[0].add_run('1. '+WordCloud.NH_negative_top3[0]+'\n')
    NHTableCells3[1].paragraphs[0].add_run('2. '+WordCloud.NH_negative_top3[1]+'\n')
    NHTableCells3[1].paragraphs[0].add_run('3. '+WordCloud.NH_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')

    # 우리은행 개인 앱 리뷰 현황
    document.add_paragraph('        - 우리은행')

    # 우리은행 표 생성
    WOORI1Table = document.add_table(3, 2, grid_t_style)

    WOORICells1 = WOORI1Table.rows[0].cells
    WOORICells1[0].paragraphs[0].add_run('긍정적 반응').bold = True
    WOORICells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORICells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    WOORI1TableCells1 = WOORI1Table.rows[0].cells
    WOORICells1[1].paragraphs[0].add_run('부정적 반응').bold = True
    WOORI1TableCells1[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1TableCells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입-우리
    WOORICell10 = WOORI1Table.cell(1, 0)
    WOORIPara10 = WOORICell10.add_paragraph()
    WOORIRun10 = WOORIPara10.add_run()
    WOORIRun10.add_picture("./wordcloud/개인고객/WOORI_WordCloud_P.png", width=Cm(7), height=Cm(5))

    WOORICell11 = WOORI1Table.cell(1, 1)
    WOORIPara11 = WOORICell11.add_paragraph()
    WOORIRun11 = WOORIPara11.add_run()
    WOORIRun11.add_picture("./wordcloud/개인고객/WOORI_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    WOORITableCells3 = WOORI1Table.rows[2].cells
    WOORITableCells3[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    WOORITableCells3[0].paragraphs[0].add_run('1. '+WordCloud.WOORI_positive_top3[0]+'\n')
    WOORITableCells3[0].paragraphs[0].add_run('2. '+WordCloud.WOORI_positive_top3[1]+'\n')
    WOORITableCells3[0].paragraphs[0].add_run('3. '+WordCloud.WOORI_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    WOORITableCells3[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    WOORITableCells3[1].paragraphs[0].add_run('1. '+WordCloud.WOORI_negative_top3[0]+'\n')
    WOORITableCells3[1].paragraphs[0].add_run('2. '+WordCloud.WOORI_negative_top3[1]+'\n')
    WOORITableCells3[1].paragraphs[0].add_run('3. '+WordCloud.WOORI_negative_top3[2]+'\n')

    document.add_paragraph('')

    # 임시 변수 !!!!!!!!!!!!!!!!!!!!!! 나중에 지우기
    #bestWord = 'UI 개선'
    resultPersonal = document.add_paragraph('타행 중 best 은행: ' +bestBankI+ '\n')
    resultPersonal.add_run('').bold = True
    resultPersonal.add_run('사용자가 가장 긍정적으로 고려한 부분은 '+bestWordI + ' 입니다.')

    for i in range(13):
        document.add_paragraph('')

    # 당행 기업 앱 리뷰 현황
    document.add_paragraph('□ 당행과 타행의 기업고객용 모바일 앱 반응 비교')
    document.add_paragraph('    ○ 당행 기업고객용 모바일 앱 (i-one bank) 사용자 반응 분석')
    IBKTableE = document.add_table(3, 2, grid_t_style)

    IBKTableECells1 = IBKTableE.rows[0].cells
    IBKTableECells1[0].paragraphs[0].add_run('긍정적 반응').bold = True
    IBKTableECells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # IBKTableECells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    IBKTableECells1 = IBKTableE.rows[0].cells
    IBKTableECells1[1].paragraphs[0].add_run('부정적 반응').bold = True
    IBKTableECells1[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # IBKTableECells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    IBKECell10 = IBKTableE.cell(1, 0)
    IBKEPara10 = IBKECell10.add_paragraph()
    IBKERun10 = IBKEPara10.add_run()
    IBKERun10.add_picture("./wordcloud/기업고객/IBK_E_WordCloud_P.png", width=Cm(7), height=Cm(5))

    IBKECell11 = IBKTableE.cell(1, 1)
    IBKEPara11 = IBKECell11.add_paragraph()
    IBKERun11 = IBKEPara11.add_run()
    IBKERun11.add_picture("./wordcloud/기업고객/IBK_E_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    IBKTableECells3 = IBKTableE.rows[2].cells
    IBKTableECells3[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    IBKTableECells3[0].paragraphs[0].add_run('1. '+WordCloud.IBK_E_positive_top3[0]+'\n')
    IBKTableECells3[0].paragraphs[0].add_run('2. '+WordCloud.IBK_E_positive_top3[1]+'\n')
    IBKTableECells3[0].paragraphs[0].add_run('3. '+WordCloud.IBK_E_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    IBKTableECells3[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    IBKTableECells3[1].paragraphs[0].add_run('1. '+WordCloud.IBK_E_negative_top3[0]+'\n')
    IBKTableECells3[1].paragraphs[0].add_run('2. '+WordCloud.IBK_E_negative_top3[1]+'\n')
    IBKTableECells3[1].paragraphs[0].add_run('3. '+WordCloud.IBK_E_negative_top3[2]+'\n')

    document.add_paragraph('')

    # 타행 기업 앱 리뷰 현황
    document.add_paragraph(' ○ 타행 기업고객용 모바일 앱 사용자 반응 분석')

    document.add_paragraph('        - 하나은행')

    # 하나은행 표 생성
    HANA1TableE = document.add_table(3, 2, grid_t_style)

    HANA1Cells1E = HANA1TableE.rows[0].cells
    HANA1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    HANA1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # HANA1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    HANA1TableCells1E = HANA1TableE.rows[0].cells
    HANA1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    HANA1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # HANA1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    HANACell10E = HANA1TableE.cell(1, 0)
    HANAPara10E = HANACell10E.add_paragraph()
    HANARun10E = HANAPara10E.add_run()
    HANARun10E.add_picture("./wordcloud/기업고객/HANA_E_WordCloud_P.png", width=Cm(7), height=Cm(5))

    HANACell11E = HANA1TableE.cell(1, 1)
    HANAPara11E = HANACell11E.add_paragraph()
    HANARun11E = HANAPara11E.add_run()
    HANARun11E.add_picture("./wordcloud/기업고객/HANA_E_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    HANATableCells3E = HANA1TableE.rows[2].cells
    HANATableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    HANATableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.HANA_E_positive_top3[0]+'\n')
    HANATableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.HANA_E_positive_top3[1]+'\n')
    HANATableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.HANA_E_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    HANATableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    HANATableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.HANA_E_negative_top3[0]+'\n')
    HANATableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.HANA_E_negative_top3[1]+'\n')
    HANATableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.HANA_E_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')

    document.add_paragraph('        - 국민은행')

    # 국민은행 표 생성
    KB1TableE = document.add_table(3, 2, grid_t_style)

    KB1Cells1E = KB1TableE.rows[0].cells
    KB1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    KB1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # KB1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    KB1TableCells1E = KB1TableE.rows[0].cells
    KB1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    KB1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # KB1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    KBCell10E = KB1TableE.cell(1, 0)
    KBPara10E = KBCell10E.add_paragraph()
    KBRun10E = KBPara10E.add_run()
    KBRun10E.add_picture("./wordcloud/기업고객/KB_E_WordCloud_P.png", width=Cm(7), height=Cm(5))

    KBCell11E = KB1TableE.cell(1, 1)
    KBPara11E = KBCell11E.add_paragraph()
    KBRun11E = KBPara11E.add_run()
    KBRun11E.add_picture("./wordcloud/기업고객/KB_E_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    KBTableCells3E = KB1TableE.rows[2].cells
    KBTableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    KBTableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.KB_E_positive_top3[0]+'\n')
    KBTableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.KB_E_positive_top3[1]+'\n')
    KBTableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.KB_E_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    KBTableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    KBTableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.KB_E_negative_top3[0]+'\n')
    KBTableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.KB_E_negative_top3[1]+'\n')
    KBTableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.KB_E_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')

    # 신한 은행 기업 앱 리뷰 현황
    document.add_paragraph('        - 신한은행')

    # 신한은행 표 생성
    SH1TableE = document.add_table(3, 2, grid_t_style)

    SH1Cells1E = SH1TableE.rows[0].cells
    SH1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    SH1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # SH1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    SH1TableCells1E = SH1TableE.rows[0].cells
    SH1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    SH1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # SH1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    SHCell10E = SH1TableE.cell(1, 0)
    SHPara10E = SHCell10E.add_paragraph()
    SHRun10E = SHPara10E.add_run()
    SHRun10E.add_picture("./wordcloud/기업고객/SHINHAN_E_WordCloud_P.png", width=Cm(7), height=Cm(5))

    SHCell11E = SH1TableE.cell(1, 1)
    SHPara11E = SHCell11E.add_paragraph()
    SHRun11E = SHPara11E.add_run()
    SHRun11E.add_picture("./wordcloud/기업고객/SHINHAN_E_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    SHTableCells3E = SH1TableE.rows[2].cells
    SHTableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    SHTableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.SHINHAN_E_positive_top3[0]+'\n')
    SHTableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.SHINHAN_E_positive_top3[1]+'\n')
    SHTableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.SHINHAN_E_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    SHTableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    SHTableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.SHINHAN_E_negative_top3[0]+'\n')
    SHTableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.SHINHAN_E_negative_top3[1]+'\n')
    SHTableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.SHINHAN_E_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')

    # 농협 은행 기업 고객 용 앱 리뷰 현황
    document.add_paragraph('        - 농협은행')

    # 농협은행 표 생성
    NH1TableE = document.add_table(3, 2, grid_t_style)

    NH1Cells1E = NH1TableE.rows[0].cells
    NH1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    NH1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # NH1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    NH1TableCells1E = SH1TableE.rows[0].cells
    NH1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    NH1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # NH1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    NHCell10E = NH1TableE.cell(1, 0)
    NHPara10E = NHCell10E.add_paragraph()
    NHRun10E = NHPara10E.add_run()
    NHRun10E.add_picture("./wordcloud/기업고객/NH_E_WordCloud_P.png", width=Cm(7), height=Cm(5))

    NHCell11E = NH1TableE.cell(1, 1)
    NHPara11E = NHCell11E.add_paragraph()
    NHRun11E = NHPara11E.add_run()
    NHRun11E.add_picture("./wordcloud/기업고객/NH_E_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    NHTableCells3E = NH1TableE.rows[2].cells
    NHTableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    NHTableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.NH_E_positive_top3[0]+'\n')
    NHTableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.NH_E_positive_top3[1]+'\n')
    NHTableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.NH_E_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    NHTableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    NHTableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.NH_E_negative_top3[0]+'\n')
    NHTableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.NH_E_negative_top3[1]+'\n')
    NHTableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.NH_E_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')

    # 우리 은행 기업 고객용 앱 현황
    document.add_paragraph('        - 우리은행')

    # 우리은행 표 생성
    WOORI1TableE = document.add_table(3, 2, grid_t_style)

    WOORI1Cells1E = WOORI1TableE.rows[0].cells
    WOORI1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    WOORI1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    WOORI1TableCells1E = WOORI1TableE.rows[0].cells
    WOORI1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    WOORI1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    WOORICell10E = WOORI1TableE.cell(1, 0)
    WOORIPara10E = WOORICell10E.add_paragraph()
    WOORIRun10E = WOORIPara10E.add_run()
    WOORIRun10E.add_picture("./wordcloud/기업고객/WOORI_E_WordCloud_P.png", width=Cm(7), height=Cm(5))

    WOORICell11E = WOORI1TableE.cell(1, 1)
    WOORIPara11E = WOORICell11E.add_paragraph()
    WOORIRun11E = WOORIPara11E.add_run()
    WOORIRun11E.add_picture("./wordcloud/기업고객/WOORI_E_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    WOORITableCells3E = WOORI1TableE.rows[2].cells
    WOORITableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    WOORITableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.WOORI_E_positive_top3[0]+'\n')
    WOORITableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.WOORI_E_positive_top3[1]+'\n')
    WOORITableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.WOORI_E_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    WOORITableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    WOORITableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.WOORI_E_negative_top3[0]+'\n')
    WOORITableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.WOORI_E_negative_top3[1]+'\n')
    WOORITableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.WOORI_E_negative_top3[2]+'\n')

    resultEnterprise = document.add_paragraph('타행 중 best 은행: ' +bestBankE+ '\n')
    resultEnterprise.add_run('').bold = True
    resultEnterprise.add_run('사용자가 가장 긍정적으로 고려한 부분은 '+bestWordE+' 입니다.')

    # 인터넷 전문 은행 앱 리뷰 현황

    document.add_paragraph('□ 인터넷 전문 은행과의 모바일 앱 사용자 반응 비교')
    document.add_paragraph('    ○ 인터넷 전문 은행')
    

    # 토스 표 생성
    document.add_paragraph('        - 토스')
    TOSS1TableE = document.add_table(3, 2, grid_t_style)

    TOSS1Cells1E = TOSS1TableE.rows[0].cells
    TOSS1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    TOSS1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    TOSS1TableCells1E = TOSS1TableE.rows[0].cells
    TOSS1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    TOSS1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    TOSSCell10E = TOSS1TableE.cell(1, 0)
    TOSSPara10E = TOSSCell10E.add_paragraph()
    TOSSRun10E = TOSSPara10E.add_run()
    TOSSRun10E.add_picture("./wordcloud/인터넷뱅크/TOSS_WordCloud_P.png", width=Cm(7), height=Cm(5))

    TOSSCell11E = TOSS1TableE.cell(1, 1)
    TOSSPara11E = TOSSCell11E.add_paragraph()
    TOSSRun11E = TOSSPara11E.add_run()
    TOSSRun11E.add_picture("./wordcloud/인터넷뱅크/TOSS_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    TOSSTableCells3E = TOSS1TableE.rows[2].cells
    TOSSTableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    TOSSTableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.TOSS_positive_top3[0]+'\n')
    TOSSTableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.TOSS_positive_top3[1]+'\n')
    TOSSTableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.TOSS_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    TOSSTableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    TOSSTableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.TOSS_negative_top3[0]+'\n')
    TOSSTableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.TOSS_negative_top3[1]+'\n')
    TOSSTableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.TOSS_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')

    document.add_paragraph('        - 카카오뱅크')

    # 카뱅 표 생성
    KAKAO1TableE = document.add_table(3, 2, grid_t_style)

    KAKAO1Cells1E = KAKAO1TableE.rows[0].cells
    KAKAO1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    KAKAO1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    KAKAO1TableCells1E = KAKAO1TableE.rows[0].cells
    KAKAO1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    KAKAO1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    KAKAOCell10E = KAKAO1TableE.cell(1, 0)
    KAKAOPara10E = KAKAOCell10E.add_paragraph()
    KAKAORun10E = KAKAOPara10E.add_run()
    KAKAORun10E.add_picture("./wordcloud/인터넷뱅크/KAKAO_WordCloud_P.png", width=Cm(7), height=Cm(5))

    KAKAOCell11E = KAKAO1TableE.cell(1, 1)
    KAKAOPara11E = KAKAOCell11E.add_paragraph()
    KAKAORun11E = KAKAOPara11E.add_run()
    KAKAORun11E.add_picture("./wordcloud/인터넷뱅크/KAKAO_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    KAKAOTableCells3E = KAKAO1TableE.rows[2].cells
    KAKAOTableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    KAKAOTableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.KAKAO_positive_top3[0]+'\n')
    KAKAOTableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.KAKAO_positive_top3[1]+'\n')
    KAKAOTableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.KAKAO_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    KAKAOTableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    KAKAOTableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.KAKAO_negative_top3[0]+'\n')
    KAKAOTableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.KAKAO_negative_top3[1]+'\n')
    KAKAOTableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.KAKAO_negative_top3[2]+'\n')

  

    # Kbank
    document.add_paragraph('        - 케이뱅크')

    # Kbank 표 생성
    Kbank1TableE = document.add_table(3, 2, grid_t_style)

    Kbank1Cells1E = Kbank1TableE.rows[0].cells
    Kbank1Cells1E[0].paragraphs[0].add_run('긍정적 반응').bold = True
    Kbank1Cells1E[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1Cells1E[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    Kbank1TableCells1E = Kbank1TableE.rows[0].cells
    Kbank1Cells1E[1].paragraphs[0].add_run('부정적 반응').bold = True
    Kbank1TableCells1E[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # WOORI1TableCells1E[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 표에 워드클라우드 삽입
    KbankCell10E = Kbank1TableE.cell(1, 0)
    KbankPara10E = KbankCell10E.add_paragraph()
    KbankRun10E = KbankPara10E.add_run()
    KbankRun10E.add_picture("./wordcloud/인터넷뱅크/Kbank_WordCloud_P.png", width=Cm(7), height=Cm(5))

    KbankCell11E = Kbank1TableE.cell(1, 1)
    KbankPara11E = KbankCell11E.add_paragraph()
    KbankRun11E = KbankPara11E.add_run()
    KbankRun11E.add_picture("./wordcloud/인터넷뱅크/Kbank_WordCloud_N.png", width=Cm(7), height=Cm(5))

    # 긍정 빈출 단어 Top3
    KbankTableCells3E = Kbank1TableE.rows[2].cells
    KbankTableCells3E[0].paragraphs[0].add_run('빈출 단어 Top3\n')
    KbankTableCells3E[0].paragraphs[0].add_run('1. '+WordCloud.KBank_positive_top3[0]+'\n')
    KbankTableCells3E[0].paragraphs[0].add_run('2. '+WordCloud.KBank_positive_top3[1]+'\n')
    KbankTableCells3E[0].paragraphs[0].add_run('3. '+WordCloud.KBank_positive_top3[2]+'\n')

    # 부정 빈출 단어 Top3
    KbankTableCells3E[1].paragraphs[0].add_run('빈출 단어 Top3\n')
    KbankTableCells3E[1].paragraphs[0].add_run('1. '+WordCloud.KBank_negative_top3[0]+'\n')
    KbankTableCells3E[1].paragraphs[0].add_run('2. '+WordCloud.KBank_negative_top3[1]+'\n')
    KbankTableCells3E[1].paragraphs[0].add_run('3. '+WordCloud.KBank_negative_top3[2]+'\n')

    document.add_paragraph('')
    document.add_paragraph('')
    
    resultInternet = document.add_paragraph('인터넷 전문 은행 중 best 은행: '+bestBankInternet+'\n')
    resultInternet.add_run('').bold = True
    resultInternet.add_run('사용자가 가장 긍정적으로 고려한 부분은 '+bestWordInternet + ' 입니다.')

    # 마지막 꼬릿말
    #document.add_paragraph('\"새로운 60년, 고객을 향한 혁신\"')

    # 문단별 정렬
    paragraph1 = document.paragraphs[0]  # 첫번째 문단
    paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph2 = document.paragraphs[3]
    paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraphLast = document.paragraphs[-1]
    # paragraphLast.font.size=Document.shared.Pt(20)
    #paragraphLast.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 마지막 문단
    # paragraphLast.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 마지막 문단

    # 파일 저장 // 마지막 단계
    document.save("report.docx")

    # 보고서를 이메일로 발송
    s = smtplib.SMTP('smtp.gmail.com', 587)  # gmail 포트번호 587
    s.starttls()  # TLS(Transport Layer Security) 보안
    s.login('IBK.ITgroup.2@gmail.com', 'czhoerpcnkfzqsdh')  # 메일을 보내는 계정
    # 메일 정보
    msg = MIMEMultipart()
    msg['From'] = 'IBK.ITgroup.2@gmail.com'
    msg['To'] = 'bethh05108@gmail.com'
    msg['Subject'] = datetime.today().strftime("%Y. %m. %d") + " I-one bank 사용자 반응 보고서입니다."
    # 메일 내용
    content = datetime.today().strftime("%Y. %m. %d") + "I-one bank 사용자 반응 보고서입니다."
    part2 = MIMEText(content, 'plain')
    msg.attach(part2)
    # 보고서 첨부
    part = MIMEBase('application', 'octet-stream')
    with open("report.docx", 'rb') as file:
        part.set_payload(file.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment', filename='report.docx')
    msg.attach(part)

    # 메일 전송
    s.sendmail("IBK.ITgroup.2@gmail.com", "bethh05108@gmail.com", msg.as_string())
    # 세션 종료
    s.quit()
